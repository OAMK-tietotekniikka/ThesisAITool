# pip install aiofiles pydantic[email] python-docx PyMuPDF pdfplumber chardet python-multipart fastapi uvicorn passlib[bcrypt] python-jose[cryptography] python-dotenv requests beautifulsoup4 pandas openai

import fitz  # PyMuPDF
import docx
import pdfplumber
import chardet
import pandas as pd
from bs4 import BeautifulSoup

import asyncio
import os, time, re, math
import uuid
import requests, random
import json
import aiohttp
from datetime import datetime
from typing import List, Optional, AsyncGenerator, Dict, Any
from enum import Enum

from fastapi import (
    FastAPI, 
    UploadFile, 
    File, 
    Form, 
    HTTPException, 
    Depends, 
    status,
    Security,
    Request
)
from fastapi.security import (
    OAuth2PasswordBearer, 
    OAuth2PasswordRequestForm,
    HTTPBearer,
    HTTPAuthorizationCredentials
)
from fastapi import HTTPException
from fastapi.responses import JSONResponse, FileResponse, HTMLResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, EmailStr, Field
from passlib.context import CryptContext
import jwt
from jwt import PyJWTError
import aiofiles

import traceback
import logging
logging.basicConfig(level=logging.DEBUG)
logging.getLogger("pdfminer").setLevel(logging.WARNING)

# Image processing imports
try:
    from PIL import Image, ImageDraw, ImageFont
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.units import inch
    IMAGE_PROCESSING_AVAILABLE = True
except ImportError:
    IMAGE_PROCESSING_AVAILABLE = False
    print("Warning: Image processing libraries not available. Thesis preview as images will be disabled.")

# Import our configuration and database
from config import config
from database import user_repo, thesis_repo, feedback_repo

# AI Provider Enum
class AIProvider(str, Enum):
    OPENAI = "openai"
    DEEPSEEK = "deepseek"
    OPENROUTER = "openrouter"

# Initialize FastAPI
app = FastAPI(
    title="ThesisAI API",
    description="API for ThesisAI application with student, supervisor, and admin roles",
    version="1.0.0",
    debug=config.DEBUG,
)

# CORS Middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount static files from client directory
app.mount("/web", StaticFiles(directory="../client"), name="web")

# Add a test route to verify static files
@app.get("/test-static")
async def test_static():
    return HTMLResponse(content="""
    <h1>Static Files Test</h1>
    <p>If you can see this, the server is running.</p>
    <p>Try these links:</p>
    <ul>
        <li><a href="/web/">/web/</a> - Client directory listing</li>
        <li><a href="/web/index.html">/web/index.html</a> - Main HTML file</li>
        <li><a href="/web/main.js">/web/main.js</a> - JavaScript file</li>
        <li><a href="/web/style.css">/web/style.css</a> - CSS file</li>
    </ul>
    """)

# Database models (using SQLite database)
class User(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    username: str
    email: EmailStr
    full_name: str
    hashed_password: str
    role: str  # "student", "supervisor", or "admin"
    disabled: bool = False
    supervisor_id: Optional[str] = None  # For students
    assigned_students: List[str] = []  # For supervisors

class Thesis(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    student_id: str
    filename: str
    filepath: str
    upload_date: datetime = Field(default_factory=datetime.now)
    status: str = "pending"  # pending, reviewed_by_ai, reviewed_by_supervisor, approved
    ai_feedback_id: Optional[str] = None
    supervisor_feedback_id: Optional[str] = None

class Feedback(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    thesis_id: str
    reviewer_id: str  # AI or supervisor ID
    content: str
    created_at: datetime = Field(default_factory=datetime.now)
    is_ai_feedback: bool

class AIRequest(BaseModel):
    thesis_id: str
    custom_instructions: str
    predefined_questions: List[str]
    provider: AIProvider = None  # Will use active provider if None
    model: Optional[str] = None

# Security
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")
security = HTTPBearer()

# Utility functions
def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password):
    return pwd_context.hash(password)

def get_user(username: str):
    return user_repo.get_user_by_username(username)

def authenticate_user(username: str, password: str):
    user_dict = get_user(username)
    if not user_dict:
        return False
    if not verify_password(password, user_dict['hashed_password']):
        return False
    return User(**user_dict)

def create_access_token(data: dict):
    to_encode = data.copy()
    encoded_jwt = jwt.encode(to_encode, config.SECRET_KEY, algorithm=config.ALGORITHM)
    return encoded_jwt

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        token = credentials.credentials
        payload = jwt.decode(token, config.SECRET_KEY, algorithms=[config.ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            raise credentials_exception
    except PyJWTError:
        raise credentials_exception
    
    user_dict = get_user(username)
    if user_dict is None:
        raise credentials_exception
    return User(**user_dict)

async def get_current_active_user(current_user: User = Depends(get_current_user)):
    if current_user.disabled:
        raise HTTPException(status_code=400, detail="Inactive user")
    return current_user

def get_student_name(student_id: str) -> str:
    student_dict = user_repo.get_user_by_id(student_id)
    if student_dict and student_dict['role'] == "student":
        return student_dict['full_name']
    else:
        raise HTTPException(status_code=404, detail="Student not found")

def check_admin(user: User):
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin privileges required")

def check_supervisor(user: User):
    if user.role != "supervisor":
        raise HTTPException(status_code=403, detail="Supervisor privileges required")

def check_student(user: User):
    if user.role != "student":
        raise HTTPException(status_code=403, detail="Student privileges required")


# Function to detect file format and extract text
def extract_text_from_file(file_path: str) -> str:
    file_ext = os.path.splitext(file_path)[1].lower()

    if file_ext == ".txt":
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        try:
            with open(file_path, "rb") as file:
                raw_data = file.read()
                detected_encoding = chardet.detect(raw_data)['encoding']
                if detected_encoding:
                    return raw_data.decode(detected_encoding)
                else:
                    return raw_data.decode('utf-8', errors='ignore')
        except Exception as e:
            print(f"Error reading text file: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Error reading text file: {str(e)}")

    elif file_ext == ".pdf":
        try:
            with pdfplumber.open(file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    text += page.extract_text()
                return text
        except Exception as e:
            print(f"Error extracting PDF text: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Error extracting PDF text: {str(e)}")

    elif file_ext == ".docx":
        try:
            doc = docx.Document(file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text
        except Exception as e:
            print(f"Error extracting DOCX text: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Error extracting DOCX text: {str(e)}")

    else:
        print("Unsupported file format")
        raise HTTPException(status_code=400, detail="Unsupported file format")

def convert_document_to_images(file_path: str, max_pages: int = 5) -> List[Dict[str, Any]]:
    """Convert document pages to images for preview"""
    if not IMAGE_PROCESSING_AVAILABLE:
        raise HTTPException(status_code=500, detail="Image processing not available")
    
    file_ext = os.path.splitext(file_path)[1].lower()
    images = []
    
    try:
        if file_ext == ".pdf":
            # Convert PDF pages to images
            pdf_document = fitz.open(file_path)
            num_pages = min(len(pdf_document), max_pages)
            
            for page_num in range(num_pages):
                page = pdf_document[page_num]
                # Render page to image with higher resolution
                mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better quality
                pix = page.get_pixmap(matrix=mat)
                
                # Convert to PIL Image
                img_data = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_data))
                
                # Resize for web display (max width 800px)
                max_width = 800
                if img.width > max_width:
                    ratio = max_width / img.width
                    new_height = int(img.height * ratio)
                    img = img.resize((max_width, new_height), Image.Resampling.LANCZOS)
                
                # Convert to base64 for web display
                img_buffer = io.BytesIO()
                img.save(img_buffer, format='PNG')
                img_base64 = base64.b64encode(img_buffer.getvalue()).decode('utf-8')
                
                images.append({
                    'page': page_num + 1,
                    'image': f"data:image/png;base64,{img_base64}",
                    'width': img.width,
                    'height': img.height
                })
            
            pdf_document.close()
            
        elif file_ext in [".doc", ".docx"]:
            # For DOC/DOCX, we'll create a simple text-based preview
            # since converting DOC/DOCX to images is complex
            try:
                doc = docx.Document(file_path)
                text_content = ""
                for para in doc.paragraphs:
                    text_content += para.text + "\n"
                
                # Create a simple text preview image
                img = create_text_preview_image(text_content[:2000], "Document Preview")  # First 2000 chars
                img_buffer = io.BytesIO()
                img.save(img_buffer, format='PNG')
                img_base64 = base64.b64encode(img_buffer.getvalue()).decode('utf-8')
                
                images.append({
                    'page': 1,
                    'image': f"data:image/png;base64,{img_base64}",
                    'width': img.width,
                    'height': img.height,
                    'text_content': text_content
                })
                
            except Exception as e:
                print(f"Error processing DOC/DOCX: {str(e)}")
                # Create a fallback image
                img = create_error_preview_image("Unable to preview document")
                img_buffer = io.BytesIO()
                img.save(img_buffer, format='PNG')
                img_base64 = base64.b64encode(img_buffer.getvalue()).decode('utf-8')
                
                images.append({
                    'page': 1,
                    'image': f"data:image/png;base64,{img_base64}",
                    'width': img.width,
                    'height': img.height
                })
        
        else:
            # For unsupported formats, create an error image
            img = create_error_preview_image("Unsupported file format")
            img_buffer = io.BytesIO()
            img.save(img_buffer, format='PNG')
            img_base64 = base64.b64encode(img_buffer.getvalue()).decode('utf-8')
            
            images.append({
                'page': 1,
                'image': f"data:image/png;base64,{img_base64}",
                'width': img.width,
                'height': img.height
            })
    
    except Exception as e:
        print(f"Error converting document to images: {str(e)}")
        # Create error image
        img = create_error_preview_image("Error processing document")
        img_buffer = io.BytesIO()
        img.save(img_buffer, format='PNG')
        img_base64 = base64.b64encode(img_buffer.getvalue()).decode('utf-8')
        
        images.append({
            'page': 1,
            'image': f"data:image/png;base64,{img_base64}",
            'width': img.width,
            'height': img.height
        })
    
    return images

def create_text_preview_image(text: str, title: str) -> Image.Image:
    """Create a preview image from text content"""
    # Create image with white background
    img = Image.new('RGB', (800, 600), color='white')
    draw = ImageDraw.Draw(img)
    
    # Try to use a default font, fallback to basic if not available
    try:
        font = ImageFont.truetype("arial.ttf", 14)
        title_font = ImageFont.truetype("arial.ttf", 18)
    except:
        font = ImageFont.load_default()
        title_font = ImageFont.load_default()
    
    # Draw title
    draw.text((20, 20), title, fill='black', font=title_font)
    
    # Draw text content with word wrapping
    lines = []
    words = text.split()
    current_line = ""
    
    for word in words:
        test_line = current_line + " " + word if current_line else word
        bbox = draw.textbbox((0, 0), test_line, font=font)
        if bbox[2] > 760:  # 800 - 40 (margins)
            if current_line:
                lines.append(current_line)
                current_line = word
            else:
                lines.append(word)
        else:
            current_line = test_line
    
    if current_line:
        lines.append(current_line)
    
    # Draw lines
    y_position = 60
    for line in lines[:30]:  # Limit to 30 lines
        draw.text((20, y_position), line, fill='black', font=font)
        y_position += 18
        if y_position > 550:
            break
    
    if len(lines) > 30:
        draw.text((20, y_position), "...", fill='gray', font=font)
    
    return img

def create_error_preview_image(message: str) -> Image.Image:
    """Create an error preview image"""
    img = Image.new('RGB', (400, 200), color='#f8f9fa')
    draw = ImageDraw.Draw(img)
    
    try:
        font = ImageFont.truetype("arial.ttf", 16)
    except:
        font = ImageFont.load_default()
    
    # Draw error message
    draw.text((20, 80), message, fill='#dc3545', font=font)
    draw.text((20, 110), "Please download the file to view its contents.", fill='#6c757d', font=font)
    
    return img

# Unified AI Interface
class UnifiedAIModel:
    def __init__(self):
        self.config = config
        self.seed = config.AI_SEED
        self.provider_config = config.get_ai_provider_config()

    def get_api_key(self, provider: AIProvider) -> Optional[str]:
        """Get API key for the specified provider"""
        provider_name = provider.value
        return self.provider_config.get(provider_name, {}).get('api_key')

    def get_model(self, provider: AIProvider, model: Optional[str] = None) -> str:
        """Get the model to use for the specified provider"""
        if model:
            return model
        provider_name = provider.value
        return self.provider_config.get(provider_name, {}).get('default_model', 'gpt-4o')

    def get_headers(self, provider: AIProvider) -> Dict[str, str]:
        """Get headers for the specified provider"""
        api_key = self.get_api_key(provider)
        if not api_key:
            raise HTTPException(status_code=500, detail=f"No API key configured for {provider}")
        
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        # Add provider-specific headers
        if provider == AIProvider.OPENROUTER:
            headers.update({
                "HTTP-Referer": "http://localhost",
                "X-Title": "ThesisAI"
            })
        
        return headers

    def get_api_url(self, provider: AIProvider) -> str:
        """Get API URL for the specified provider"""
        provider_name = provider.value
        return self.provider_config.get(provider_name, {}).get('api_url', '')

    async def make_request(self, provider: AIProvider, messages: List[Dict[str, str]], 
                          model: Optional[str] = None, stream: bool = False) -> Dict[str, Any]:
        """Make a request to the specified AI provider"""
        api_key = self.get_api_key(provider)
        if not api_key:
            raise HTTPException(status_code=500, detail=f"No API key configured for {provider}")
        
        model_name = self.get_model(provider, model)
        headers = self.get_headers(provider)
        api_url = self.get_api_url(provider)
        
        payload = {
            "model": model_name,
            "messages": messages,
            "stream": stream,
        }
        
        # Add provider-specific parameters
        if provider == AIProvider.OPENROUTER:
            payload["seed"] = self.seed
        
        try:
            response = requests.post(api_url, headers=headers, json=payload, timeout=60)
            response.raise_for_status()
            return response.json()
        except requests.exceptions.RequestException as e:
            print(f"❌ Error with {provider}: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Error with {provider}: {str(e)}")

    async def make_streaming_request(self, provider: AIProvider, messages: List[Dict[str, str]], 
                                   model: Optional[str] = None, pacing_delay: float = 0.01) -> AsyncGenerator[str, None]:
        """Make a streaming request to the specified AI provider with improved UX"""
        api_key = self.get_api_key(provider)
        if not api_key:
            print(f"🔄 Using fallback mode for {provider} - no API key available")
            yield f"data: {json.dumps({'type': 'status', 'content': f'[FALLBACK MODE] {provider.value.upper()} Analysis'})}\n\n"
            yield f"data: {json.dumps({'type': 'content', 'content': f'This is a simulated response since no API key is configured for {provider}.'})}\n\n"
            yield f"data: {json.dumps({'type': 'complete'})}\n\n"
            return
        
        model_name = self.get_model(provider, model)
        headers = self.get_headers(provider)
        api_url = self.get_api_url(provider)
        
        payload = {
            "model": model_name,
            "messages": messages,
            "stream": True,
        }
        
        if provider == AIProvider.OPENROUTER:
            payload["seed"] = self.seed
        
        try:
            # Send initial status
            yield f"data: {json.dumps({'type': 'status', 'content': f'Connecting to {provider.value.upper()}...'})}\n\n"
            
            async with aiohttp.ClientSession() as session:
                async with session.post(api_url, headers=headers, json=payload, timeout=aiohttp.ClientTimeout(total=120)) as response:
                    if response.status != 200:
                        error_text = await response.text()
                        print(f"❌ HTTP Error: {response.status} - {error_text}")
                        yield f"data: {json.dumps({'type': 'error', 'content': f'Failed to get AI feedback from {provider}: {response.status}'})}\n\n"
                        return
                    
                    # Send connected status
                    yield f"data: {json.dumps({'type': 'status', 'content': f'Connected to {provider.value.upper()}. Generating response...'})}\n\n"
                    
                    buffer = ""
                    async for line in response.content:
                        line_str = line.decode('utf-8').strip()
                        if line_str.startswith('data: '):
                            data_content = line_str[6:]
                            if data_content == '[DONE]':
                                yield f"data: {json.dumps({'type': 'complete'})}\n\n"
                                break
                            try:
                                json_data = json.loads(data_content)
                                if 'choices' in json_data and len(json_data['choices']) > 0:
                                    delta = json_data['choices'][0].get('delta', {})
                                    if 'content' in delta:
                                        content = delta['content']
                                        buffer += content
                                        
                                        # Send content in chunks for better UX
                                        if len(buffer) >= 10 or '\n' in buffer:  # Send every 10 chars or on newline
                                            yield f"data: {json.dumps({'type': 'content', 'content': buffer})}\n\n"
                                            buffer = ""
                                            await asyncio.sleep(pacing_delay)  # Control pacing
                                        
                            except json.JSONDecodeError:
                                continue
                    
                    # Send any remaining buffer
                    if buffer:
                        yield f"data: {json.dumps({'type': 'content', 'content': buffer})}\n\n"
                        
        except asyncio.TimeoutError:
            print(f"❌ Timeout with {provider}")
            yield f"data: {json.dumps({'type': 'error', 'content': f'Request timed out for {provider}'})}\n\n"
        except Exception as e:
            print(f"❌ Error with {provider} streaming: {str(e)}")
            yield f"data: {json.dumps({'type': 'error', 'content': f'Error with {provider}: {str(e)}'})}\n\n"



    async def analyze_thesis_stream(self, file_path: str, custom_instructions: str, 
                                  predefined_questions: List[str], provider: AIProvider = None, 
                                  model: Optional[str] = None) -> AsyncGenerator[str, None]:
        """Stream thesis analysis using the specified AI provider with enhanced UX"""
        if provider is None:
            provider = AIProvider(config.get_active_provider())
            
        try:
            thesis_content = extract_text_from_file(file_path)
        except Exception as e:
            print(f"Error reading thesis file: {str(e)}")
            yield f"data: {json.dumps({'type': 'error', 'content': f'Error reading thesis file: {str(e)}'})}\n\n"
            return
        
        prompt = f"Analyze the following thesis content: {thesis_content}\nPlease answer the following questions:\n"
        for question in predefined_questions:
            prompt += f"- {question}\n"
        prompt += "\nIMPORTANT: Provide direct answers to the questions above. Do NOT ask any follow-up questions. Do NOT ask for clarification. Simply provide your analysis and recommendations based on the content provided. Thank you!"
        
        messages = [{"role": "user", "content": prompt}]
        
        if custom_instructions:
            messages.insert(0, {"role": "system", "content": custom_instructions + "\n\nCRITICAL INSTRUCTION: You must NOT ask any follow-up questions. Provide direct analysis and feedback only."})
        else:
            messages.insert(0, {"role": "system", "content": "You are a thesis evaluation assistant. Provide direct analysis and feedback. Do NOT ask follow-up questions or request clarification. Give comprehensive answers based on the information provided."})
        
        async for chunk in self.make_streaming_request(provider, messages, model):
            yield chunk

    async def grade_formatting_style(self, file_path: str, provider: AIProvider = None, 
                                   model: Optional[str] = None) -> AsyncGenerator[str, None]:
        """Stream formatting style analysis using the specified AI provider"""
        if provider is None:
            provider = AIProvider(config.get_active_provider())
            
        try:
            thesis_content = extract_text_from_file(file_path)
        except Exception as e:
            print(f"Error reading thesis file: {str(e)}")
            yield f"data: {json.dumps({'type': 'error', 'content': f'Error reading thesis file: {str(e)}'})}\n\n"
            return

        with open('oamk_ref_guidelines.txt', 'r', encoding='utf8') as f:
            guidelines = f.read()

        prompt = f"""
        Analyze the following thesis content and provide detailed analysis of FORMATTING STYLE.
        
        Thesis Content:
        {thesis_content}
        
        WHAT to detect:
        - Detect incorrect reference style, for example number references, and suggest referecing to be used. You MUST indicate where the problem is (page number or chapter).
        - Detect incorrect reference style of tables and figures.
        - Detect incorrect formatting of citations and references.

        STRICTLY follow these guidelines:
        [Start of guidelines]
        {guidelines}.
        [end of guidelines]
        
        Provide specific examples from the thesis to support your analysis.
        
        IMPORTANT: Provide direct analysis and evaluation. Do NOT ask any follow-up questions or request clarification.
        """
        
        messages = [{"role": "user", "content": prompt}]
        messages.insert(0, {"role": "system", "content": "You are a thesis evaluation assistant. Provide direct analysis and feedback. Do NOT ask follow-up questions or request clarification. Give comprehensive answers based on the information provided."})
        
        async for chunk in self.make_streaming_request(provider, messages, model):
            yield chunk

    async def grade_purpose_objectives(self, file_path: str, provider: AIProvider = None, 
                                     model: Optional[str] = None) -> AsyncGenerator[str, None]:
        """Stream purpose and objectives analysis using the specified AI provider"""
        if provider is None:
            provider = AIProvider(config.get_active_provider())
            
        try:
            thesis_content = extract_text_from_file(file_path)
        except Exception as e:
            print(f"Error reading thesis file: {str(e)}")
            yield f"data: {json.dumps({'type': 'error', 'content': f'Error reading thesis file: {str(e)}'})}\n\n"
            return
        
        prompt = f"""
        Analyze the following thesis content and provide detailed analysis of PURPOSE AND OBJECTIVES.
        
        Thesis Content:
        {thesis_content}
        
        Please evaluate the purpose and objectives strictly following this grading scale:
        
        Excellent (5): Purposes and objectives are well-grounded in theory and practice, and are directed toward the application of professional development results.
        
        Good (4-3): Purposes and objectives are directed toward the development of the professional field.
        
        Satisfactory (2-1): The thesis has a basic objective.
        
        Fail (0)/Unfinished: Objectives are vague or not in accordance with the approved plan.
        
        Please provide:
        1. A grade (0-5) based on the above scale
        2. Detailed justification for the grade
        3. Specific examples from the thesis to support your evaluation
        4. Recommendations for improvement if applicable
        
        IMPORTANT: Provide direct analysis and evaluation. Do NOT ask any follow-up questions or request clarification.
        """
        
        messages = [{"role": "user", "content": prompt}]
        messages.insert(0, {"role": "system", "content": "You are a thesis evaluation assistant. Provide direct analysis and feedback. Do NOT ask follow-up questions or request clarification. Give comprehensive answers based on the information provided."})
        
        async for chunk in self.make_streaming_request(provider, messages, model):
            yield chunk

    async def grade_theoretical_foundation(self, file_path: str, provider: AIProvider = None, 
                                         model: Optional[str] = None) -> AsyncGenerator[str, None]:
        """Stream theoretical foundation analysis using the specified AI provider"""
        if provider is None:
            provider = AIProvider(config.get_active_provider())
            
        try:
            thesis_content = extract_text_from_file(file_path)
        except Exception as e:
            print(f"Error reading thesis file: {str(e)}")
            yield f"data: {json.dumps({'type': 'error', 'content': f'Error reading thesis file: {str(e)}'})}\n\n"
            return
        
        prompt = f"""
        Analyze the following thesis content and provide detailed analysis of THEORETICAL FOUNDATION.
        
        Thesis Content:
        {thesis_content}
        
        Please evaluate the theoretical foundation strictly following this grading scale:
        
        Excellent (5): The theoretical foundation conveys the author's own, critical and creative thinking. It is carefully considered, topical and purposeful in terms of the nature of the work. A sufficient amount of key scientific/artistic research and specialist knowledge has been used for the theoretical foundation.
        
        Good (4-3): The thesis has a theoretical foundation and is based on versatile industry sources.
        
        Satisfactory (2-1): The thesis has a theoretical foundation and is based on industry sources.
        
        Fail (0)/Unfinished: The theoretical foundation is noticeably limited and selected uncritically.
        
        Please provide:
        1. A grade (0-5) based on the above scale
        2. Detailed justification for the grade
        3. Specific examples from the thesis to support your evaluation
        4. Recommendations for improvement if applicable
        
        IMPORTANT: Provide direct analysis and evaluation. Do NOT ask any follow-up questions or request clarification.
        """
        
        messages = [{"role": "user", "content": prompt}]
        messages.insert(0, {"role": "system", "content": "You are a thesis evaluation assistant. Provide direct analysis and feedback. Do NOT ask follow-up questions or request clarification. Give comprehensive answers based on the information provided."})
        
        async for chunk in self.make_streaming_request(provider, messages, model):
            yield chunk

    async def grade_professional_connection(self, file_path: str, provider: AIProvider = None, 
                                          model: Optional[str] = None) -> AsyncGenerator[str, None]:
        """Stream professional connection analysis using the specified AI provider"""
        if provider is None:
            provider = AIProvider(config.get_active_provider())
            
        try:
            thesis_content = extract_text_from_file(file_path)
        except Exception as e:
            print(f"Error reading thesis file: {str(e)}")
            yield f"data: {json.dumps({'type': 'error', 'content': f'Error reading thesis file: {str(e)}'})}\n\n"
            return
        
        prompt = f"""
        Analyze the following thesis content and provide detailed analysis of CONNECTION OF SUBJECT TO PROFESSIONAL FIELD AND EXPERTISE.
        
        Thesis Content:
        {thesis_content}
        
        Please evaluate the connection to professional field and expertise strictly following this grading scale:
        
        Excellent (5): The subject has a well-argued connection to the professional field and it plays an important role in developing the student's expertise. The subject is valuable for practical activity and important for working life and its development. The subject is of current interest, new, creative, demanding.
        
        Good (4-3): The subject is clear connection to the professional field and it is related to the student's professional development. The subject is valuable and well-reasoned from a worklife perspective. The subject is of current interest and typical of the field.
        
        Satisfactory (2-1): The subject is related to the development of the industry and the student's professional growth. The subject is useful for the working life/client. The subject is ordinary.
        
        Fail (0)/Unfinished: The subject has no connection to the professional field.
        
        Please provide:
        1. A grade (0-5) based on the above scale
        2. Detailed justification for the grade
        3. Specific examples from the thesis to support your evaluation
        4. Recommendations for improvement if applicable
        
        IMPORTANT: Provide direct analysis and evaluation. Do NOT ask any follow-up questions or request clarification.
        """
        
        messages = [{"role": "user", "content": prompt}]
        messages.insert(0, {"role": "system", "content": "You are a thesis evaluation assistant. Provide direct analysis and feedback. Do NOT ask follow-up questions or request clarification. Give comprehensive answers based on the information provided."})
        
        async for chunk in self.make_streaming_request(provider, messages, model):
            yield chunk

    async def grade_development_task(self, file_path: str, provider: AIProvider = None, 
                                   model: Optional[str] = None) -> AsyncGenerator[str, None]:
        """Stream development task analysis using the specified AI provider"""
        if provider is None:
            provider = AIProvider(config.get_active_provider())
            
        try:
            thesis_content = extract_text_from_file(file_path)
        except Exception as e:
            print(f"Error reading thesis file: {str(e)}")
            yield f"data: {json.dumps({'type': 'error', 'content': f'Error reading thesis file: {str(e)}'})}\n\n"
            return
        
        prompt = f"""
        Analyze the following thesis content and provide detailed analysis of DEVELOPMENT/RESEARCH TASK AND ITS DEFINITION.
        
        Thesis Content:
        {thesis_content}
        
        Please evaluate the development/research task and its definition strictly following this grading scale:
        
        Excellent (5): The development/research task and its definition are described clearly and justified.
        
        Good (4-3): The development/research task and its definition are well-argued.
        
        Satisfactory (2-1): The development/research task is understood.
        
        Fail (0)/Unfinished: The development/research task has not been defined.
        
        Please provide:
        1. A grade (0-5) based on the above scale
        2. Detailed justification for the grade
        3. Specific examples from the thesis to support your evaluation
        4. Recommendations for improvement if applicable
        
        IMPORTANT: Provide direct analysis and evaluation. Do NOT ask any follow-up questions or request clarification.
        """
        
        messages = [{"role": "user", "content": prompt}]
        messages.insert(0, {"role": "system", "content": "You are a thesis evaluation assistant. Provide direct analysis and feedback. Do NOT ask follow-up questions or request clarification. Give comprehensive answers based on the information provided."})
        
        async for chunk in self.make_streaming_request(provider, messages, model):
            yield chunk

    async def grade_conclusions_proposals(self, file_path: str, provider: AIProvider = None, 
                                        model: Optional[str] = None) -> AsyncGenerator[str, None]:
        """Stream conclusions and proposals analysis using the specified AI provider"""
        if provider is None:
            provider = AIProvider(config.get_active_provider())
            
        try:
            thesis_content = extract_text_from_file(file_path)
        except Exception as e:
            print(f"Error reading thesis file: {str(e)}")
            yield f"data: {json.dumps({'type': 'error', 'content': f'Error reading thesis file: {str(e)}'})}\n\n"
            return
        
        prompt = f"""
        Analyze the following thesis content and provide detailed analysis of CONCLUSIONS/DEVELOPMENT PROPOSALS.
        
        Thesis Content:
        {thesis_content}
        
        Please evaluate the conclusions/development proposals strictly following this grading scale:
        
        Excellent (5): Conclusions/development proposals reflect the results themselves compared to the research data and expertise.
        
        Good (4-3): Conclusions/development proposals are normal and appropriate.
        
        Satisfactory (2-1): Basic conclusions/recommendations are given.
        
        Fail (0)/Unfinished: No conclusions/recommendations.
        
        Please provide:
        1. A grade (0-5) based on the above scale
        2. Detailed justification for the grade
        3. Specific examples from the thesis to support your evaluation
        4. Recommendations for improvement if applicable
        
        IMPORTANT: Provide direct analysis and evaluation. Do NOT ask any follow-up questions or request clarification.
        """
        
        messages = [{"role": "user", "content": prompt}]
        messages.insert(0, {"role": "system", "content": "You are a thesis evaluation assistant. Provide direct analysis and feedback. Do NOT ask follow-up questions or request clarification. Give comprehensive answers based on the information provided."})
        
        async for chunk in self.make_streaming_request(provider, messages, model):
            yield chunk

    async def grade_material_methodology(self, file_path: str, provider: AIProvider = None, 
                                       model: Optional[str] = None) -> AsyncGenerator[str, None]:
        """Stream material and methodology analysis using the specified AI provider"""
        if provider is None:
            provider = AIProvider(config.get_active_provider())
            
        try:
            thesis_content = extract_text_from_file(file_path)
        except Exception as e:
            print(f"Error reading thesis file: {str(e)}")
            yield f"data: {json.dumps({'type': 'error', 'content': f'Error reading thesis file: {str(e)}'})}\n\n"
            return
        
        prompt = f"""
        Analyze the following thesis content and provide detailed analysis of MATERIAL AND METHODOLOGICAL CHOICES.
        
        Thesis Content:
        {thesis_content}
        
        Please evaluate the material and methodological choices strictly following this grading scale:
        
        Excellent (5): The material is diverse from the viewpoint of the objective of the work. The acquisition of material and work methods are well-founded and their use is well-controlled.
        
        Good (4-3): The material is comprehensive. The acquisition of material and work methods are well-founded.
        
        Satisfactory (2-1): The material is sufficient. The acquisition of material and work methods are purposeful, and they have been described.
        
        Fail (0)/Unfinished: The material is insufficient. The acquisition of material and work methods have not been described.
        
        Please provide:
        1. A grade (0-5) based on the above scale
        2. Detailed justification for the grade
        3. Specific examples from the thesis to support your evaluation
        4. Recommendations for improvement if applicable
        
        IMPORTANT: Provide direct analysis and evaluation. Do NOT ask any follow-up questions or request clarification.
        """
        
        messages = [{"role": "user", "content": prompt}]
        messages.insert(0, {"role": "system", "content": "You are a thesis evaluation assistant. Provide direct analysis and feedback. Do NOT ask follow-up questions or request clarification. Give comprehensive answers based on the information provided."})
        
        async for chunk in self.make_streaming_request(provider, messages, model):
            yield chunk

    async def grade_treatment_analysis(self, file_path: str, provider: AIProvider = None, 
                                     model: Optional[str] = None) -> AsyncGenerator[str, None]:
        """Stream treatment and analysis analysis using the specified AI provider"""
        if provider is None:
            provider = AIProvider(config.get_active_provider())
            
        try:
            thesis_content = extract_text_from_file(file_path)
        except Exception as e:
            print(f"Error reading thesis file: {str(e)}")
            yield f"data: {json.dumps({'type': 'error', 'content': f'Error reading thesis file: {str(e)}'})}\n\n"
            return
        
        prompt = f"""
        Analyze the following thesis content and provide detailed analysis of TREATMENT AND ANALYSIS OF MATERIAL.
        
        Thesis Content:
        {thesis_content}
        
        Please evaluate the treatment and analysis of material strictly following this grading scale:
        
        Excellent (5): The material is treated in a controlled manner and analysis is proficient. It shows a creative and systematic approach.
        
        Good (4-3): The treatment and analysis of material illustrates the author's familiarity with the subject.
        
        Satisfactory (2-1): The treatment and analysis of material is adequate.
        
        Fail (0)/Unfinished: The treatment and analysis of material is inconsistent and inconsistent.
        
        Please provide:
        1. A grade (0-5) based on the above scale
        2. Detailed justification for the grade
        3. Specific examples from the thesis to support your evaluation
        4. Recommendations for improvement if applicable
        
        IMPORTANT: Provide direct analysis and evaluation. Do NOT ask any follow-up questions or request clarification.
        """
        
        messages = [{"role": "user", "content": prompt}]
        messages.insert(0, {"role": "system", "content": "You are a thesis evaluation assistant. Provide direct analysis and feedback. Do NOT ask follow-up questions or request clarification. Give comprehensive answers based on the information provided."})
        
        async for chunk in self.make_streaming_request(provider, messages, model):
            yield chunk

    async def grade_results_product(self, file_path: str, provider: AIProvider = None, 
                                  model: Optional[str] = None) -> AsyncGenerator[str, None]:
        """Stream results and product analysis using the specified AI provider"""
        if provider is None:
            provider = AIProvider(config.get_active_provider())
            
        try:
            thesis_content = extract_text_from_file(file_path)
        except Exception as e:
            print(f"Error reading thesis file: {str(e)}")
            yield f"data: {json.dumps({'type': 'error', 'content': f'Error reading thesis file: {str(e)}'})}\n\n"
            return
        
        prompt = f"""
        Analyze the following thesis content and provide detailed analysis of RESULTS/PRODUCT.
        
        Thesis Content:
        {thesis_content}
        
        Please evaluate the results/product strictly following this grading scale:
        
        Excellent (5): The results/product are new creations and original, the application of results has been proven and significance assessed.
        
        Good (4-3): The objectives set for the work have been justified. The achieved results/product can be applied to the development of the industry.
        
        Satisfactory (2-1): The objectives set for the work have been reached.
        
        Failed (0)/Unfinished: The objectives set for the work have not been reached. The results have been wrongly interpreted.
        
        Please provide:
        1. A grade (0-5) based on the above scale
        2. Detailed justification for the grade
        3. Specific examples from the thesis to support your evaluation
        4. Recommendations for improvement if applicable
        
        IMPORTANT: Provide direct analysis and evaluation. Do NOT ask any follow-up questions or request clarification.
        """
        
        messages = [{"role": "user", "content": prompt}]
        messages.insert(0, {"role": "system", "content": "You are a thesis evaluation assistant. Provide direct analysis and feedback. Do NOT ask follow-up questions or request clarification. Give comprehensive answers based on the information provided."})
        
        async for chunk in self.make_streaming_request(provider, messages, model):
            yield chunk

# Initialize the unified AI model
ai_model = UnifiedAIModel()

# Routes
@app.post("/token")
async def login(form_data: OAuth2PasswordRequestForm = Depends()):
    print("form_data:", form_data);
    user = authenticate_user(form_data.username, form_data.password)
    if not user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    access_token = create_access_token(
        data={"sub": user.username}
    )
    return {"access_token": access_token, "token_type": "bearer"}

@app.post("/register")
async def register(
    username: str = Form(...),
    email: EmailStr = Form(...),
    full_name: str = Form(...),
    password: str = Form(...),
    role: str = Form(...),
    supervisor_id: Optional[str] = Form(None),
):
    if user_repo.get_user_by_username(username):
        raise HTTPException(status_code=400, detail="Username already registered")
    
    if role not in ["student", "supervisor", "admin"]:
        raise HTTPException(status_code=400, detail="Invalid role")
    
    hashed_password = get_password_hash(password)
    user = User(
        username=username,
        email=email,
        full_name=full_name,
        hashed_password=hashed_password,
        role=role,
        supervisor_id=supervisor_id
    )
    
    user_repo.add_user(user.dict())
    
    if role == "student" and supervisor_id:
        supervisor = user_repo.get_user_by_username(supervisor_id)
        if supervisor:
            user_repo.add_assigned_student(supervisor['id'], user.id)
    
    return {"message": "User created successfully", "user_id": user.id}

@app.post("/upload-thesis")
async def upload_thesis(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_active_user)
):
    check_student(current_user)
    
    print(f"📤 Uploading thesis for user: {current_user.username} (ID: {current_user.id})")
    print(f"📄 File name: {file.filename}")
    
    file_ext = os.path.splitext(file.filename)[1]
    unique_filename = f"{current_user.id}_{datetime.now().strftime('%Y%m%d%H%M%S')}{file_ext}"
    file_path = os.path.join(config.UPLOAD_DIR, unique_filename)
    
    try:
        async with aiofiles.open(file_path, 'wb') as f:
            while content := await file.read(1024):
                await f.write(content)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error saving file: {str(e)}")
    
    thesis = Thesis(
        student_id=current_user.id,
        filename=file.filename,
        filepath=file_path,
    )
    
    print(f"📝 Created thesis with ID: {thesis.id}")
    thesis_repo.add_thesis(thesis.dict())
    
    return {"message": "Thesis uploaded successfully", "thesis_id": thesis.id}

@app.get("/", response_class=HTMLResponse)
async def root():
    try:
        with open("./_baocao/index.html", encoding='utf8') as f:
            return HTMLResponse(content=f.read())
    except FileNotFoundError:
        html_content = """
        <!DOCTYPE html>
        <html>
        <head>
            <title>ThesisAI API</title>
            <meta charset="utf-8">
        </head>
        <body>
            <h1>ThesisAI API</h1>
            <p>The API is running successfully!</p>
            <p>Check the <a href="/docs">API documentation</a> for available endpoints.</p>
        </body>
        </html>
        """
        return HTMLResponse(content=html_content)

async def stream_ai_feedback(thesis_id: str, custom_instructions: str, predefined_questions: List[str], 
                           provider: AIProvider = None, model: Optional[str] = None) -> AsyncGenerator[str, None]:
    """Stream AI feedback for a thesis using the specified provider with enhanced UX and meaningful chunk buffering"""
    thesis = thesis_repo.get_thesis_by_id(thesis_id)
    if not thesis:
        print(f"❌ Thesis not found: {thesis_id}")
        yield f"data: {json.dumps({'type': 'error', 'content': 'Thesis not found'})}\n\n"
        yield f"data: {json.dumps({'type': 'complete'})}\n\n"
        return
    
    print(f"✅ Starting AI feedback for thesis: {thesis_id}")
    print(f"📄 File path: {thesis['filepath']}")
    print(f"🤖 Provider: {provider or 'active'}")
    print(f"🤖 Model: {model or 'default'}")
    
    try:
        if not os.path.exists(thesis['filepath']):
            print(f"❌ Thesis file not found: {thesis['filepath']}")
            yield f"data: {json.dumps({'type': 'error', 'content': 'Thesis file not found'})}\n\n"
            yield f"data: {json.dumps({'type': 'complete'})}\n\n"
            return
        
        # Send initial progress
        yield f"data: {json.dumps({'type': 'progress', 'content': 'Starting thesis analysis...', 'step': 1, 'total': 3})}\n\n"
        
        # Step 1: Thesis Analysis
        print("🔄 Starting thesis analysis...")
        buffer = ""
        try:
            async for chunk in ai_model.analyze_thesis_stream(thesis['filepath'], custom_instructions, predefined_questions, provider, model):
                # Parse the chunk to extract structured data
                if chunk.startswith('data: '):
                    try:
                        data = json.loads(chunk[6:])
                        if data.get('type') == 'content':
                            # Buffer the content
                            buffer += data.get('content', '')
                            
                            # Send meaningful chunks (sentences, paragraphs, or after certain length)
                            if len(buffer) >= 50 or '\n\n' in buffer or buffer.endswith(('.', '!', '?')):
                                yield f"data: {json.dumps({'type': 'content', 'content': buffer})}\n\n"
                                buffer = ""
                                await asyncio.sleep(0.1)  # Small delay for better UX
                        elif data.get('type') == 'error':
                            yield chunk
                            return
                        elif data.get('type') == 'complete':
                            # Send any remaining buffer
                            if buffer:
                                yield f"data: {json.dumps({'type': 'content', 'content': buffer})}\n\n"
                            break
                    except json.JSONDecodeError:
                        # Handle legacy format
                        yield chunk
                else:
                    # Handle non-JSON chunks
                    buffer += chunk
                    if len(buffer) >= 50 or '\n\n' in buffer or buffer.endswith(('.', '!', '?')):
                        yield f"data: {json.dumps({'type': 'content', 'content': buffer})}\n\n"
                        buffer = ""
                        await asyncio.sleep(0.1)
        except (asyncio.CancelledError, GeneratorExit):
            print(f"🔴 [STOP STREAM] Client disconnected during thesis analysis for thesis_id: {thesis_id}")
            print(f"🔴 [STOP STREAM] Stopped AI model to save resources.")  # AI models auto-stop on disconnects
            return
        
        # Send any remaining buffer from analysis
        if buffer:
            yield f"data: {json.dumps({'type': 'content', 'content': buffer})}\n\n"
            
        yield f"data: {json.dumps({'type': 'progress', 'content': 'Thesis analysis completed. Starting objective grading...', 'step': 2, 'total': 3})}\n\n"
        yield f"data: {json.dumps({'type': 'section', 'content': 'GRADING PURPOSES AND OBJECTIVES'})}\n\n"
        
        # Step 2: Objective Grading
        print("🔄 Starting objective grading...")
        buffer = ""
        try:
            async for chunk in ai_model.grade_purpose_objectives(thesis['filepath'], provider, model):
                if chunk.startswith('data: '):
                    try:
                        data = json.loads(chunk[6:])
                        if data.get('type') == 'content':
                            buffer += data.get('content', '')
                            
                            if len(buffer) >= 50 or '\n\n' in buffer or buffer.endswith(('.', '!', '?')):
                                yield f"data: {json.dumps({'type': 'content', 'content': buffer})}\n\n"
                                buffer = ""
                                await asyncio.sleep(0.1)
                        elif data.get('type') == 'error':
                            yield chunk
                            return
                        elif data.get('type') == 'complete':
                            if buffer:
                                yield f"data: {json.dumps({'type': 'content', 'content': buffer})}\n\n"
                            break
                    except json.JSONDecodeError:
                        yield chunk
                else:
                    buffer += chunk
                    if len(buffer) >= 50 or '\n\n' in buffer or buffer.endswith(('.', '!', '?')):
                        yield f"data: {json.dumps({'type': 'content', 'content': buffer})}\n\n"
                        buffer = ""
                        await asyncio.sleep(0.1)
        except (asyncio.CancelledError, GeneratorExit):
            print(f"🔴 [STOP STREAM] Client disconnected during objective grading for thesis_id: {thesis_id}")
            print(f"🔴 [STOP STREAM] Stopped AI model to save resources.")  # AI models auto-stop on disconnects
            return
        
        # Send any remaining buffer from objective grading
        if buffer:
            yield f"data: {json.dumps({'type': 'content', 'content': buffer})}\n\n"
                
        yield f"data: {json.dumps({'type': 'progress', 'content': 'Objective grading completed. Starting theoretical foundation grading...', 'step': 3, 'total': 3})}\n\n"
        yield f"data: {json.dumps({'type': 'section', 'content': 'GRADING THEORETICAL FOUNDATION'})}\n\n"
        
        # Step 3: Theoretical Foundation Grading
        print("🔄 Starting theoretical foundation grading...")
        buffer = ""
        try:
            async for chunk in ai_model.grade_theoretical_foundation(thesis['filepath'], provider, model):
                if chunk.startswith('data: '):
                    try:
                        data = json.loads(chunk[6:])
                        if data.get('type') == 'content':
                            buffer += data.get('content', '')
                            
                            if len(buffer) >= 50 or '\n\n' in buffer or buffer.endswith(('.', '!', '?')):
                                yield f"data: {json.dumps({'type': 'content', 'content': buffer})}\n\n"
                                buffer = ""
                                await asyncio.sleep(0.1)
                        elif data.get('type') == 'error':
                            yield chunk
                            return
                        elif data.get('type') == 'complete':
                            if buffer:
                                yield f"data: {json.dumps({'type': 'content', 'content': buffer})}\n\n"
                            break
                    except json.JSONDecodeError:
                        yield chunk
                else:
                    buffer += chunk
                    if len(buffer) >= 50 or '\n\n' in buffer or buffer.endswith(('.', '!', '?')):
                        yield f"data: {json.dumps({'type': 'content', 'content': buffer})}\n\n"
                        buffer = ""
                        await asyncio.sleep(0.1)
        except (asyncio.CancelledError, GeneratorExit):
            print(f"🔴 [STOP STREAM] Client disconnected during theoretical foundation grading for thesis_id: {thesis_id}")
            print(f"🔴 [STOP STREAM] Stopped AI model to save resources.")  # AI models auto-stop on disconnects
            return
        
        # Send any remaining buffer from theoretical foundation grading
        if buffer:
            yield f"data: {json.dumps({'type': 'content', 'content': buffer})}\n\n"
            
        print("✅ AI feedback streaming completed successfully")
        yield f"data: {json.dumps({'type': 'progress', 'content': 'Analysis completed successfully!', 'step': 3, 'total': 3})}\n\n"
        yield f"data: {json.dumps({'type': 'complete'})}\n\n"
            
    except (asyncio.CancelledError, GeneratorExit):
        print(f"🔴 [STOP STREAM] Client disconnected during streaming for thesis_id: {thesis_id}")
        print(f"🔴 [STOP STREAM] Stopped AI model to save resources.")  # AI models auto-stop on disconnects
        return
    except Exception as e:
        print("❌ Full traceback:\n", traceback.format_exc())
        print(f"❌ Error in stream_ai_feedback: {str(e)}")
        yield f"data: {json.dumps({'type': 'error', 'content': f'AI service error: {str(e)}'})}\n\n"
        yield f"data: {json.dumps({'type': 'complete'})}\n\n"

@app.post("/request-ai-feedback")
async def request_ai_feedback(
    thesis_id: str,
    custom_instructions: str = Form(""),
    predefined_questions: List[str] = Form([]),
    selected_options: str = Form(""),
    current_user: User = Depends(get_current_active_user)
):
    """Request AI feedback for a thesis with streaming response"""
    
    # Validate thesis access
    thesis = thesis_repo.get_thesis_by_id(thesis_id)
    if not thesis:
        raise HTTPException(status_code=404, detail="Thesis not found")
    
    if current_user.role == "student" and thesis['student_id'] != current_user.id:
        raise HTTPException(status_code=403, detail="Not your thesis")
    
    if current_user.role == "supervisor" and thesis['student_id'] not in current_user.assigned_students:
        raise HTTPException(status_code=403, detail="Not your assigned student")
    
    # Parse selected options if provided
    selected_options_list = []
    if selected_options:
        try:
            selected_options_list = json.loads(selected_options)
        except json.JSONDecodeError:
            pass
    
    # Use the new grade functions if selected_options are provided
    if selected_options_list:
        return StreamingResponse(
            stream_ai_feedback_with_grades(thesis_id, selected_options_list),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Headers": "Cache-Control"
            }
        )
    
    # Use predefined questions if provided
    if predefined_questions:
        return StreamingResponse(
            stream_ai_feedback(thesis_id, custom_instructions, predefined_questions),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Headers": "Cache-Control"
            }
        )
    
    # Fallback to default questions if none provided
    predefined_questions = [
        "What are the strengths of this thesis?",
        "What areas need improvement?",
        "How well is the methodology implemented?",
        "Are references properly formatted?",
        "How strong is the theoretical foundation?"
    ]
    
    return StreamingResponse(
        stream_ai_feedback(thesis_id, custom_instructions, predefined_questions),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Headers": "Cache-Control"
        }
    )

@app.post("/save-ai-feedback")
async def save_ai_feedback(
    thesis_id: str,
    feedback_content: str = Form(...),
    current_user: User = Depends(get_current_active_user)
):
    thesis = thesis_repo.get_thesis_by_id(thesis_id)
    if not thesis:
        raise HTTPException(status_code=404, detail="Thesis not found")
    
    if thesis['student_id'] != current_user.id:
        raise HTTPException(status_code=403, detail="Not your thesis")
    
    feedback = Feedback(
        thesis_id=thesis_id,
        reviewer_id="ai_system",
        content=feedback_content,
        is_ai_feedback=True
    )
    feedback_repo.add_feedback(feedback.dict())
    
    thesis_repo.update_thesis_ai_feedback(thesis_id, feedback.id)
    thesis_repo.update_thesis_status(thesis_id, "reviewed_by_ai")
    
    ai_response_path = os.path.join(config.AI_RESPONSES_DIR, f"{thesis_id}_ai_response.txt")
    async with aiofiles.open(ai_response_path, 'w', encoding='utf-8') as f:
        await f.write(feedback_content)

    return {"message": "Feedback saved successfully", "feedback_id": feedback.id}

@app.post("/submit-supervisor-feedback")
async def submit_supervisor_feedback(
    thesis_id: str,
    feedback_content: str = Form(...),
    current_user: User = Depends(get_current_active_user)
):
    check_supervisor(current_user)
    
    thesis = thesis_repo.get_thesis_by_id(thesis_id)
    if not thesis:
        raise HTTPException(status_code=404, detail="Thesis not found")
    
    if thesis['student_id'] not in current_user.assigned_students:
        raise HTTPException(status_code=403, detail="Not your assigned student")
    
    feedback = Feedback(
        thesis_id=thesis_id,
        reviewer_id=current_user.id,
        content=feedback_content,
        is_ai_feedback=False
    )
    feedback_repo.add_feedback(feedback.dict())
    
    thesis_repo.update_thesis_supervisor_feedback(thesis_id, feedback.id)
    thesis_repo.update_thesis_status(thesis_id, "reviewed_by_supervisor")
    
    feedback_path = os.path.join(config.FEEDBACK_DIR, f"{thesis_id}_supervisor_feedback.txt")
    async with aiofiles.open(feedback_path, 'w') as f:
        await f.write(feedback_content)
    
    return {"message": "Feedback submitted successfully", "feedback_id": feedback.id}

@app.get("/get-supervisor-feedback/{thesis_id}")
async def get_supervisor_feedback(
    thesis_id: str,
    current_user: User = Depends(get_current_active_user)
):
    thesis = thesis_repo.get_thesis_by_id(thesis_id)
    if not thesis:
        raise HTTPException(status_code=404, detail="Thesis not found")
    
    if current_user.role == "student" and thesis['student_id'] != current_user.id:
        raise HTTPException(status_code=403, detail="Not your thesis")
    
    if current_user.role == "supervisor" and thesis['student_id'] not in current_user.assigned_students:
        raise HTTPException(status_code=403, detail="Not your assigned student")
    
    if not thesis['supervisor_feedback_id']:
        raise HTTPException(status_code=404, detail="No supervisor feedback available")
    
    feedback = feedback_repo.get_feedback_by_id(thesis['supervisor_feedback_id'])
    if not feedback:
        raise HTTPException(status_code=404, detail="Feedback not found")
    
    return feedback

@app.get("/supervisor-feedback")
async def get_all_supervisor_feedback(current_user: User = Depends(get_current_active_user)):
    """Get all supervisor feedback for the current user"""
    if current_user.role == "student":
        # Get all theses for this student that have supervisor feedback
        theses = thesis_repo.get_theses_by_student_id(current_user.id)
        feedback_list = []
        
        for thesis in theses:
            if thesis.get('supervisor_feedback_id'):
                feedback = feedback_repo.get_feedback_by_id(thesis['supervisor_feedback_id'])
                if feedback:
                    # Get supervisor name
                    supervisor = user_repo.get_user_by_id(feedback['reviewer_id'])
                    supervisor_name = supervisor['full_name'] if supervisor else 'Unknown Supervisor'
                    
                    feedback_list.append({
                        'thesis_id': thesis['id'],
                        'thesis_title': thesis['filename'],
                        'feedback_text': feedback['content'],
                        'feedback_date': feedback['created_at'],
                        'supervisor_name': supervisor_name
                    })
        
        return feedback_list
    elif current_user.role == "supervisor":
        # Get all theses for students assigned to this supervisor that have supervisor feedback
        theses = thesis_repo.get_theses_by_supervisor(current_user.username)
        feedback_list = []
        
        for thesis in theses:
            if thesis.get('supervisor_feedback_id'):
                feedback = feedback_repo.get_feedback_by_id(thesis['supervisor_feedback_id'])
                if feedback and feedback['reviewer_id'] == current_user.id:
                    # Get student name
                    student = user_repo.get_user_by_id(thesis['student_id'])
                    student_name = student['full_name'] if student else 'Unknown Student'
                    
                    feedback_list.append({
                        'thesis_id': thesis['id'],
                        'thesis_title': thesis['filename'],
                        'feedback_text': feedback['content'],
                        'feedback_date': feedback['created_at'],
                        'student_name': student_name
                    })
        
        return feedback_list
    else:
        # Admin can see all supervisor feedback
        all_theses = thesis_repo.get_all_theses()
        feedback_list = []
        
        for thesis in all_theses:
            if thesis.get('supervisor_feedback_id'):
                feedback = feedback_repo.get_feedback_by_id(thesis['supervisor_feedback_id'])
                if feedback:
                    # Get supervisor and student names
                    supervisor = user_repo.get_user_by_id(feedback['reviewer_id'])
                    student = user_repo.get_user_by_id(thesis['student_id'])
                    supervisor_name = supervisor['full_name'] if supervisor else 'Unknown Supervisor'
                    student_name = student['full_name'] if student else 'Unknown Student'
                    
                    feedback_list.append({
                        'thesis_id': thesis['id'],
                        'thesis_title': thesis['filename'],
                        'feedback_text': feedback['content'],
                        'feedback_date': feedback['created_at'],
                        'supervisor_name': supervisor_name,
                        'student_name': student_name
                    })
        
        return feedback_list

@app.post("/assign-supervisor")
async def assign_supervisor(
    student_username: str = Form(...),
    supervisor_username: str = Form(...),
    current_user: User = Depends(get_current_active_user)
):
    check_admin(current_user)
    
    student = user_repo.get_user_by_username(student_username)
    if not student or student['role'] != "student":
        raise HTTPException(status_code=404, detail="Student not found")
    
    supervisor = user_repo.get_user_by_username(supervisor_username)
    if not supervisor or supervisor['role'] != "supervisor":
        raise HTTPException(status_code=404, detail="Supervisor not found")
    
    # Use the repository method for assignment
    success = user_repo.assign_supervisor(student_username, supervisor_username)
    if not success:
        raise HTTPException(status_code=500, detail="Failed to assign supervisor")
    
    return {"message": f"Supervisor {supervisor_username} assigned to student {student_username}"}

@app.get("/my-theses")
async def get_my_theses(current_user: User = Depends(get_current_active_user)):
    print(f"🔍 Getting theses for user: {current_user.username} (ID: {current_user.id})")
    print(f"🔍 User role: {current_user.role}")
    
    if current_user.role == "student":
        theses = thesis_repo.get_theses_by_student_id(current_user.id)
        print(f"🔍 Found {len(theses)} theses for student")
        # Add student name for student's own theses
        for thesis in theses:
            thesis['student_name'] = current_user.full_name
    elif current_user.role == "supervisor":
        theses = thesis_repo.get_theses_by_supervisor(current_user.username)
        print(f"🔍 Found {len(theses)} theses for supervisor")
        # Add student names for supervisor's assigned students
        for thesis in theses:
            student = user_repo.get_user_by_id(thesis['student_id'])
            thesis['student_name'] = student['full_name'] if student else 'Unknown Student'
    else:  # admin
        theses = thesis_repo.get_all_theses()
        print(f"🔍 Found {len(theses)} theses for admin")
        # Add student names for all theses
        for thesis in theses:
            student = user_repo.get_user_by_id(thesis['student_id'])
            thesis['student_name'] = student['full_name'] if student else 'Unknown Student'
    
    print(f"🔍 Returning theses: {[{'id': t['id'], 'filename': t['filename'], 'student_id': t['student_id'], 'student_name': t.get('student_name', 'N/A')} for t in theses]}")
    
    return theses

@app.get("/download-thesis/{thesis_id}")
async def download_thesis(
    thesis_id: str,
    current_user: User = Depends(get_current_active_user)
):
    thesis = thesis_repo.get_thesis_by_id(thesis_id)
    if not thesis:
        raise HTTPException(status_code=404, detail="Thesis not found")
    
    if current_user.role == "student" and thesis['student_id'] != current_user.id:
        raise HTTPException(status_code=403, detail="Not your thesis")
    
    if current_user.role == "supervisor" and thesis['student_id'] not in current_user.assigned_students:
        raise HTTPException(status_code=403, detail="Not your assigned student")
    
    if not os.path.exists(thesis['filepath']):
        raise HTTPException(status_code=404, detail="Thesis file not found")
    
    return FileResponse(
        thesis['filepath'],
        filename=thesis['filename'],
        media_type="application/octet-stream"
    )

@app.get("/theses-to-review")
async def get_theses_to_review(current_user: User = Depends(get_current_active_user)):
    check_supervisor(current_user)
    theses_to_review = []
    
    # Get theses assigned to this supervisor
    supervisor_theses = thesis_repo.get_theses_by_supervisor(current_user.username)
    
    for thesis in supervisor_theses:
        if thesis['status'] in ["reviewed_by_ai", "pending"]:
            student = user_repo.get_user_by_id(thesis['student_id'])
            if student:
                thesis_data = {
                    "student_name": student['full_name'],
                    "filename": thesis['filename'],
                    "upload_date": thesis['upload_date'],
                    "status": thesis['status']
                }
                theses_to_review.append(thesis_data)

    return theses_to_review

@app.get("/users")
async def get_users(current_user: User = Depends(get_current_active_user)):
    check_admin(current_user)
    return list(user_repo.get_all_users())

@app.get("/me")
async def read_users_me(current_user: User = Depends(get_current_active_user)):
    return current_user

@app.get("/supervisors")
async def get_supervisors(current_user: User = Depends(get_current_active_user)):
    supervisors = [u for u in user_repo.get_all_users() if u['role'] == "supervisor"]
    return supervisors

@app.get("/students")
async def get_students(current_user: User = Depends(get_current_active_user)):
    if current_user.role == "supervisor":
        # Get students assigned to this supervisor
        supervisor = user_repo.get_user_by_username(current_user.username)
        if supervisor and supervisor['assigned_students']:
            students = []
            for student_username in supervisor['assigned_students']:
                student = user_repo.get_user_by_username(student_username)
                if student:
                    students.append(student)
        else:
            students = []
    else:
        check_admin(current_user)
        students = user_repo.get_users_by_role("student")
    return students

@app.get("/supervisor-assignments")
async def get_supervisor_assignments(current_user: User = Depends(get_current_active_user)):
    check_admin(current_user)
    assignments = []
    students = user_repo.get_users_by_role("student")
    
    for student in students:
        supervisor_name = None
        if student['supervisor_id']:
            supervisor = user_repo.get_user_by_username(student['supervisor_id'])
            if supervisor:
                supervisor_name = supervisor['full_name']
        assignments.append({
            "student_id": student['username'],
            "student_name": student['full_name'],
            "supervisor_name": supervisor_name
        })
    return assignments

@app.get("/all-theses")
async def get_all_theses(current_user: User = Depends(get_current_active_user)):
    check_admin(current_user)
    
    theses_with_student_names = []
    theses = thesis_repo.get_all_theses()
    
    for thesis in theses:
        student = user_repo.get_user_by_id(thesis['student_id'])
        if student and student['role'] == "student":
            thesis_dict = thesis.copy()
            thesis_dict["student_name"] = student['full_name']
            theses_with_student_names.append(thesis_dict)
    
    return theses_with_student_names

@app.get("/ai-providers")
async def get_ai_providers():
    """Get available AI providers and their configuration"""
    return config.get_available_providers()

@app.get("/config/status")
async def get_config_status():
    """Get configuration status including JWT and AI provider status"""
    jwt_valid = config.validate_jwt_config()
    ai_status = config.validate_ai_config()
    
    return {
        "jwt": {
            "valid": jwt_valid,
            "algorithm": config.ALGORITHM,
            "expire_minutes": config.ACCESS_TOKEN_EXPIRE_MINUTES
        },
        "ai_providers": ai_status,
        "active_provider": config.get_active_provider(),
        "server": {
            "host": config.HOST,
            "port": config.PORT,
            "debug": config.DEBUG
        },
        "directories": {
            "upload_dir": config.UPLOAD_DIR,
            "feedback_dir": config.FEEDBACK_DIR,
            "ai_responses_dir": config.AI_RESPONSES_DIR
        }
    }

@app.get("/streaming-config")
async def get_streaming_config():
    """Get streaming configuration for client-side optimization"""
    return {
        "pacing_delay": 0.01,  # seconds between chunks
        "buffer_size": 10,      # characters per chunk
        "timeout": 120,         # seconds
        "retry_attempts": 3,
        "supported_types": [
            "content",      # Regular content chunks
            "status",       # Status updates
            "progress",     # Progress indicators
            "section",      # Section headers
            "error",        # Error messages
            "complete"      # Stream completion
        ]
    }

@app.post("/request-ai-feedback-enhanced")
async def request_ai_feedback_enhanced(
    thesis_id: str,
    custom_instructions: str = Form("Please review this thesis and provide feedback"),
    predefined_questions: List[str] = Form(["What are the strengths?", "What areas need improvement?"]),
    provider: AIProvider = Form(None),
    model: Optional[str] = Form(None),
    pacing_delay: float = Form(0.01),
    current_user: User = Depends(get_current_active_user)
):
    """Enhanced AI feedback endpoint with configurable pacing and better error handling"""
    print(f"🔍 Enhanced AI feedback request for thesis_id: {thesis_id}")
    print(f"🔍 Current user: {current_user.username} (ID: {current_user.id})")
    print(f"🔍 Provider: {provider or 'active'}")
    print(f"🔍 Model: {model}")
    print(f"🔍 Pacing delay: {pacing_delay}")
    
    thesis = thesis_repo.get_thesis_by_id(thesis_id)
    if not thesis:
        print(f"❌ Thesis not found in database. Available theses: {list(thesis_repo.get_all_theses())}")
        raise HTTPException(status_code=404, detail="Thesis not found")
    
    if thesis['student_id'] != current_user.id:
        print(f"❌ Thesis belongs to student {thesis['student_id']}, but current user is {current_user.id}")
        raise HTTPException(status_code=403, detail="Not your thesis")
    
    return StreamingResponse(
        stream_ai_feedback_enhanced(thesis_id, custom_instructions, predefined_questions, provider, model, pacing_delay),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Headers": "Cache-Control",
            "X-Streaming-Version": "2.0"
        }
    )

async def stream_ai_feedback_enhanced(thesis_id: str, custom_instructions: str, predefined_questions: List[str], 
                                     provider: AIProvider = None, model: Optional[str] = None, 
                                     pacing_delay: float = 0.01) -> AsyncGenerator[str, None]:
    """Enhanced streaming function with better pacing and error recovery"""
    thesis = thesis_repo.get_thesis_by_id(thesis_id)
    if not thesis:
        yield f"data: {json.dumps({'type': 'error', 'content': 'Thesis not found'})}\n\n"
        yield f"data: {json.dumps({'type': 'complete'})}\n\n"
        return
    
    try:
        if not os.path.exists(thesis['filepath']):
            yield f"data: {json.dumps({'type': 'error', 'content': 'Thesis file not found'})}\n\n"
            yield f"data: {json.dumps({'type': 'complete'})}\n\n"
            return
        
        # Send initial status with metadata
        yield f"data: {json.dumps({
            'type': 'status', 
            'content': 'Starting analysis...',
            'metadata': {
                'thesis_id': thesis_id,
                'provider': provider.value if provider else 'active',
                'model': model or 'default',
                'pacing_delay': pacing_delay
            }
        })}\n\n"
        
        # Step 1: Thesis Analysis
        yield f"data: {json.dumps({'type': 'progress', 'content': 'Analyzing thesis content...', 'step': 1, 'total': 3})}\n\n"
        
        try:
            async for chunk in ai_model.analyze_thesis_stream(thesis['filepath'], custom_instructions, predefined_questions, provider, model):
                if chunk.startswith('data: '):
                    try:
                        data = json.loads(chunk[6:])
                        if data.get('type') == 'content':
                            yield chunk
                            await asyncio.sleep(pacing_delay)  # Apply pacing
                        elif data.get('type') == 'error':
                            yield chunk
                            return
                        elif data.get('type') == 'complete':
                            break
                    except json.JSONDecodeError:
                        yield chunk
                else:
                    yield chunk
                    await asyncio.sleep(pacing_delay)  # Apply pacing to legacy format
        except (asyncio.CancelledError, GeneratorExit):
            print(f"🔴 [STOP STREAM] Client disconnected during thesis analysis for thesis_id: {thesis_id}")
            print(f"🔴 [STOP STREAM] Stopped AI model to save resources.")  # AI models auto-stop on disconnects
            return
        
        # Step 2: Objective Grading
        yield f"data: {json.dumps({'type': 'progress', 'content': 'Grading objectives...', 'step': 2, 'total': 3})}\n\n"
        yield f"data: {json.dumps({'type': 'section', 'content': 'GRADING PURPOSES AND OBJECTIVES'})}\n\n"
        
        try:
            async for chunk in ai_model.grade_purpose_objectives(thesis['filepath'], provider, model):
                if chunk.startswith('data: '):
                    try:
                        data = json.loads(chunk[6:])
                        if data.get('type') == 'content':
                            yield chunk
                            await asyncio.sleep(pacing_delay)
                        elif data.get('type') == 'error':
                            yield chunk
                            return
                        elif data.get('type') == 'complete':
                            break
                    except json.JSONDecodeError:
                        yield chunk
                        await asyncio.sleep(pacing_delay)
                else:
                    yield chunk
                    await asyncio.sleep(pacing_delay)
        except (asyncio.CancelledError, GeneratorExit):
            print(f"🔴 [STOP STREAM] Client disconnected during objective grading for thesis_id: {thesis_id}")
            print(f"🔴 [STOP STREAM] Stopped AI model to save resources.")  # AI models auto-stop on disconnects
            return
        
        # Step 3: Theoretical Foundation Grading
        yield f"data: {json.dumps({'type': 'progress', 'content': 'Grading theoretical foundation...', 'step': 3, 'total': 3})}\n\n"
        yield f"data: {json.dumps({'type': 'section', 'content': 'GRADING THEORETICAL FOUNDATION'})}\n\n"
        
        try:
            async for chunk in ai_model.grade_theoretical_foundation(thesis['filepath'], provider, model):
                if chunk.startswith('data: '):
                    try:
                        data = json.loads(chunk[6:])
                        if data.get('type') == 'content':
                            yield chunk
                            await asyncio.sleep(pacing_delay)
                        elif data.get('type') == 'error':
                            yield chunk
                            return
                        elif data.get('type') == 'complete':
                            break
                    except json.JSONDecodeError:
                        yield chunk
                        await asyncio.sleep(pacing_delay)
                else:
                    yield chunk
                    await asyncio.sleep(pacing_delay)
        except (asyncio.CancelledError, GeneratorExit):
            print(f"🔴 [STOP STREAM] Client disconnected during theoretical foundation grading for thesis_id: {thesis_id}")
            print(f"🔴 [STOP STREAM] Stopped AI model to save resources.")  # AI models auto-stop on disconnects
            return
        
        yield f"data: {json.dumps({'type': 'progress', 'content': 'Analysis completed successfully!', 'step': 3, 'total': 3})}\n\n"
        yield f"data: {json.dumps({'type': 'complete'})}\n\n"
            
    except (asyncio.CancelledError, GeneratorExit):
        print(f"🔴 [STOP STREAM] Client disconnected during streaming for thesis_id: {thesis_id}")
        print(f"🔴 [STOP STREAM] Stopped AI model to save resources.")  # AI models auto-stop on disconnects
        return
    except Exception as e:
        print(f"❌ Error in enhanced streaming: {str(e)}")
        yield f"data: {json.dumps({'type': 'error', 'content': f'Streaming error: {str(e)}'})}\n\n"
        yield f"data: {json.dumps({'type': 'complete'})}\n\n"

@app.get("/simple-test-page")
async def simple_test_page():
    """Serve the simple streaming test page"""
    try:
        with open("./server/simple_test.html", encoding='utf8') as f:
            return HTMLResponse(content=f.read())
    except FileNotFoundError:
        return HTMLResponse(content="<h1>Simple test page not found</h1>")

@app.get("/debug-streaming-page")
async def debug_streaming_page():
    """Serve the debug streaming tool"""
    try:
        with open("./server/debug_streaming.html", encoding='utf8') as f:
            return HTMLResponse(content=f.read())
    except FileNotFoundError:
        return HTMLResponse(content="<h1>Debug tool not found</h1>")

@app.get("/test-streaming-page")
async def test_streaming_page():
    """Serve the test streaming page"""
    try:
        with open("./server/test_streaming.html", encoding='utf8') as f:
            return HTMLResponse(content=f.read())
    except FileNotFoundError:
        return HTMLResponse(content="<h1>Test page not found</h1>")

@app.get("/test-streaming")
async def test_streaming():
    """Test endpoint to verify streaming format"""
    async def test_stream():
        yield f"data: {json.dumps({'type': 'status', 'content': 'Starting test...'})}\n\n"
        await asyncio.sleep(0.5)
        
        yield f"data: {json.dumps({'type': 'progress', 'content': 'Step 1/3', 'step': 1, 'total': 3})}\n\n"
        await asyncio.sleep(0.5)
        
        yield f"data: {json.dumps({'type': 'content', 'content': 'This is a test message. '})}\n\n"
        await asyncio.sleep(0.5)
        
        yield f"data: {json.dumps({'type': 'content', 'content': 'It should display properly. '})}\n\n"
        await asyncio.sleep(0.5)
        
        yield f"data: {json.dumps({'type': 'content', 'content': 'With proper formatting.'})}\n\n"
        await asyncio.sleep(0.5)
        
        yield f"data: {json.dumps({'type': 'complete'})}\n\n"
    
    return StreamingResponse(
        test_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Headers": "Cache-Control"
        }
    )

@app.get("/test-ai-feedback")
async def test_ai_feedback():
    """Test endpoint for AI feedback without authentication"""
    
    async def test_stream():
        yield f"data: {json.dumps({'type': 'progress', 'content': 'Starting test analysis...', 'step': 1, 'total': 2})}\n\n"
        
        # Simulate some content
        test_content = """### Test Feedback Analysis

This is a test response to verify that the streaming endpoint is working correctly.

#### Strengths
- The system is properly configured
- Streaming is functional
- JSON parsing works correctly

#### Areas for Improvement
- Add more comprehensive error handling
- Implement better progress tracking
- Enhance the user interface

#### Recommendations
1. Continue testing the system
2. Monitor performance
3. Gather user feedback

This test confirms that the AI feedback system is operational."""
        
        # Stream the content in chunks
        words = test_content.split()
        for i, word in enumerate(words):
            yield f"data: {json.dumps({'type': 'content', 'content': word + ' '})}\n\n"
            await asyncio.sleep(0.1)  # Simulate processing time
        
        yield f"data: {json.dumps({'type': 'progress', 'content': 'Test completed successfully!', 'step': 2, 'total': 2})}\n\n"
        yield f"data: {json.dumps({'type': 'complete'})}\n\n"
    
    return StreamingResponse(
        test_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Headers": "*",
        }
    )

@app.get("/extract-thesis-text/{thesis_id}")
async def extract_thesis_text(
    thesis_id: str,
    current_user: User = Depends(get_current_active_user)
):
    """Extract text content from thesis file for preview"""
    thesis = thesis_repo.get_thesis_by_id(thesis_id)
    if not thesis:
        raise HTTPException(status_code=404, detail="Thesis not found")
    
    if current_user.role == "student" and thesis['student_id'] != current_user.id:
        raise HTTPException(status_code=403, detail="Not your thesis")
    
    if current_user.role == "supervisor" and thesis['student_id'] not in current_user.assigned_students:
        raise HTTPException(status_code=403, detail="Not your assigned student")
    
    if not os.path.exists(thesis['filepath']):
        raise HTTPException(status_code=404, detail="Thesis file not found")
    
    try:
        text_content = extract_text_from_file(thesis['filepath'])
        return {"text": text_content}
    except Exception as e:
        print(f"Error extracting text from thesis: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error extracting text: {str(e)}")

@app.get("/ai-feedback-options")
async def get_ai_feedback_options():
    """Get available AI feedback options for dynamic checkbox generation"""
    return {
        "options": [
            {
                "id": "formatting_style",
                "label": "Formatting style",
                "description": "Check formatting, structure, and presentation quality",
                "enabled": True,
                "default": True
            },
            {
                "id": "purpose_objectives",
                "label": "Purpose and objectives",
                "description": "Evaluate clarity and grounding of purposes and objectives",
                "enabled": True,
                "default": True
            },
            {
                "id": "theoretical_foundation",
                "label": "Theoretical foundation",
                "description": "Assess theoretical framework and critical thinking",
                "enabled": True,
                "default": True
            },
            {
                "id": "professional_connection",
                "label": "Connection of subject to professional field and expertise",
                "description": "Evaluate relevance to professional development and working life",
                "enabled": True,
                "default": True
            },
            {
                "id": "development_task",
                "label": "Development/research task and its definition",
                "description": "Assess clarity and justification of research/development tasks",
                "enabled": True,
                "default": True
            },
            {
                "id": "conclusions_proposals",
                "label": "Conclusions/development proposals",
                "description": "Evaluate quality of conclusions and development proposals",
                "enabled": True,
                "default": True
            },
            {
                "id": "material_methodology",
                "label": "Material and methodological choices",
                "description": "Assess diversity and foundation of materials and methods",
                "enabled": True,
                "default": True
            },
            {
                "id": "treatment_analysis",
                "label": "Treatment and analysis of material",
                "description": "Evaluate controlled treatment and proficient analysis",
                "enabled": True,
                "default": True
            },
            {
                "id": "results_product",
                "label": "Results/Product",
                "description": "Assess originality and application of results",
                "enabled": True,
                "default": True
            }
        ]
    }

@app.get("/thesis-preview-images/{thesis_id}")
async def get_thesis_preview_images(
    thesis_id: str,
    current_user: User = Depends(get_current_active_user)
):
    """Get preview images for a thesis document"""
    thesis = thesis_repo.get_thesis_by_id(thesis_id)
    if not thesis:
        raise HTTPException(status_code=404, detail="Thesis not found")
    
    if current_user.role == "student" and thesis['student_id'] != current_user.id:
        raise HTTPException(status_code=403, detail="Not your thesis")
    
    if current_user.role == "supervisor" and thesis['student_id'] not in current_user.assigned_students:
        raise HTTPException(status_code=403, detail="Not your assigned student")
    
    if not os.path.exists(thesis['filepath']):
        raise HTTPException(status_code=404, detail="Thesis file not found")
    
    try:
        images = convert_document_to_images(thesis['filepath'], max_pages=5)
        return {
            "thesis_id": thesis_id,
            "filename": thesis['filename'],
            "images": images
        }
    except Exception as e:
        print(f"Error generating preview images: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating preview images: {str(e)}")

# Add new streaming function after the existing stream_ai_feedback function
async def stream_ai_feedback_with_grades(thesis_id: str, selected_options: List[str], 
                                        provider: AIProvider = None, model: Optional[str] = None) -> AsyncGenerator[str, None]:
    """Stream AI feedback using the new grade functions based on selected options"""
    thesis = thesis_repo.get_thesis_by_id(thesis_id)
    if not thesis:
        print(f"❌ Thesis not found: {thesis_id}")
        yield f"data: {json.dumps({'type': 'error', 'content': 'Thesis not found'})}\n\n"
        yield f"data: {json.dumps({'type': 'complete'})}\n\n"
        return
    
    print(f"✅ Starting AI feedback with grades for thesis: {thesis_id}")
    print(f"📄 File path: {thesis['filepath']}")
    print(f"🤖 Provider: {provider or 'active'}")
    print(f"🤖 Model: {model or 'default'}")
    print(f"📋 Selected options: {selected_options}")
    
    try:
        if not os.path.exists(thesis['filepath']):
            print(f"❌ Thesis file not found: {thesis['filepath']}")
            yield f"data: {json.dumps({'type': 'error', 'content': 'Thesis file not found'})}\n\n"
            yield f"data: {json.dumps({'type': 'complete'})}\n\n"
            return
        
        # Map of option IDs to grade functions
        grade_functions = {
            'formatting_style': ai_model.grade_formatting_style,
            'purpose_objectives': ai_model.grade_purpose_objectives,
            'theoretical_foundation': ai_model.grade_theoretical_foundation,
            'professional_connection': ai_model.grade_professional_connection,
            'development_task': ai_model.grade_development_task,
            'conclusions_proposals': ai_model.grade_conclusions_proposals,
            'material_methodology': ai_model.grade_material_methodology,
            'treatment_analysis': ai_model.grade_treatment_analysis,
            'results_product': ai_model.grade_results_product
        }
        
        # Option titles for section headers
        option_titles = {
            'formatting_style': 'FORMATTING STYLE',
            'purpose_objectives': 'PURPOSE AND OBJECTIVES',
            'theoretical_foundation': 'THEORETICAL FOUNDATION',
            'professional_connection': 'CONNECTION OF SUBJECT TO PROFESSIONAL FIELD AND EXPERTISE',
            'development_task': 'DEVELOPMENT/RESEARCH TASK AND ITS DEFINITION',
            'conclusions_proposals': 'CONCLUSIONS/DEVELOPMENT PROPOSALS',
            'material_methodology': 'MATERIAL AND METHODOLOGICAL CHOICES',
            'treatment_analysis': 'TREATMENT AND ANALYSIS OF MATERIAL',
            'results_product': 'RESULTS/PRODUCT'
        }
        
        total_options = len(selected_options)
        
        for i, option in enumerate(selected_options, 1):
            if option not in grade_functions:
                print(f"⚠️ Unknown option: {option}")
                continue
                
            print(f"🔄 Processing option {i}/{total_options}: {option}")
            
            # Send progress update
            yield f"data: {json.dumps({'type': 'progress', 'content': f'Analyzing {option_titles.get(option, option)}...', 'step': i, 'total': total_options})}\n\n"
            
            # Send section header
            yield f"data: {json.dumps({'type': 'section', 'content': option_titles.get(option, option.upper().replace('_', ' '))})}\n\n"
            
            # Get the grade function
            grade_func = grade_functions[option]
            
            # Stream the grade analysis
            buffer = ""
            try:
                async for chunk in grade_func(thesis['filepath'], provider, model):
                    if chunk.startswith('data: '):
                        try:
                            data = json.loads(chunk[6:])
                            if data.get('type') == 'content':
                                buffer += data.get('content', '')
                                
                                if len(buffer) >= 50 or '\n\n' in buffer or buffer.endswith(('.', '!', '?')):
                                    yield f"data: {json.dumps({'type': 'content', 'content': buffer})}\n\n"
                                    buffer = ""
                                    await asyncio.sleep(0.1)
                            elif data.get('type') == 'error':
                                yield chunk
                                return
                            elif data.get('type') == 'complete':
                                if buffer:
                                    yield f"data: {json.dumps({'type': 'content', 'content': buffer})}\n\n"
                                break
                        except json.JSONDecodeError:
                            yield chunk
                    else:
                        buffer += chunk
                        if len(buffer) >= 50 or '\n\n' in buffer or buffer.endswith(('.', '!', '?')):
                            yield f"data: {json.dumps({'type': 'content', 'content': buffer})}\n\n"
                            buffer = ""
                            await asyncio.sleep(0.1)
            except (asyncio.CancelledError, GeneratorExit):
                print(f"🔴 [STOP STREAM] Client disconnected during {option} analysis for thesis_id: {thesis_id}")
                print(f"🔴 [STOP STREAM] Stopped AI model to save resources.")  # AI models auto-stop on disconnects
                return
            
            # Send any remaining buffer
            if buffer:
                yield f"data: {json.dumps({'type': 'content', 'content': buffer})}\n\n"
        
        print("✅ AI feedback with grades streaming completed successfully")
        yield f"data: {json.dumps({'type': 'progress', 'content': 'Analysis completed successfully!', 'step': total_options, 'total': total_options})}\n\n"
        yield f"data: {json.dumps({'type': 'complete'})}\n\n"
            
    except (asyncio.CancelledError, GeneratorExit):
        print(f"🔴 [STOP STREAM] Client disconnected during streaming for thesis_id: {thesis_id}")
        print(f"🔴 [STOP STREAM] Stopped AI model to save resources.")  # AI models auto-stop on disconnects
        return
    except Exception as e:
        print("❌ Full traceback:\n", traceback.format_exc())
        print(f"❌ Error in stream_ai_feedback_with_grades: {str(e)}")
        yield f"data: {json.dumps({'type': 'error', 'content': f'AI service error: {str(e)}'})}\n\n"
        yield f"data: {json.dumps({'type': 'complete'})}\n\n"

if __name__ == "__main__":
    import uvicorn
    
    # Print configuration status on startup
    print("🚀 Starting ThesisAI Tool...")
    print(f"📁 Upload directory: {config.UPLOAD_DIR}")
    print(f"📁 Feedback directory: {config.FEEDBACK_DIR}")
    print(f"📁 AI responses directory: {config.AI_RESPONSES_DIR}")
    print(f"🤖 Active AI provider: {config.get_active_provider()}")
    
    # Check JWT configuration
    if not config.validate_jwt_config():
        print("⚠️  JWT configuration warning - using default secret key")
    
    # Check AI provider configuration
    ai_status = config.validate_ai_config()
    available_providers = [k for k, v in ai_status.items() if v]
    if available_providers:
        print(f"✅ Available AI providers: {', '.join(available_providers)}")
    else:
        print("⚠️  No AI providers configured - using fallback mode")
    
    uvicorn.run(app, host=config.HOST, port=config.PORT, reload=config.DEBUG) 